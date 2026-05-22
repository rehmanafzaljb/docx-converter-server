# from fastapi import FastAPI, UploadFile, File, HTTPException
# from fastapi.responses import Response
# import subprocess
# import tempfile
# import os
# import shutil

# app = FastAPI()

# @app.get("/health")
# def health():
#     return {"status": "ok"}

# @app.post("/convert/word-to-pdf")
# async def word_to_pdf(file: UploadFile = File(...)):
#     if not file.filename.endswith(('.docx', '.doc')):
#         raise HTTPException(
#             status_code=400,
#             detail="Only .docx and .doc files are supported"
#         )

#     tmp_dir = tempfile.mkdtemp()
#     try:
#         input_path = os.path.join(tmp_dir, file.filename)
#         with open(input_path, "wb") as f:
#             content = await file.read()
#             f.write(content)

#         result = subprocess.run([
#             "libreoffice",
#             "--headless",
#             "--convert-to", "pdf",
#             "--outdir", tmp_dir,
#             input_path
#         ], capture_output=True, text=True, timeout=60)

#         if result.returncode != 0:
#             raise HTTPException(
#                 status_code=500,
#                 detail=f"Conversion failed: {result.stderr}"
#             )

#         pdf_filename = os.path.splitext(file.filename)[0] + ".pdf"
#         pdf_path = os.path.join(tmp_dir, pdf_filename)

#         if not os.path.exists(pdf_path):
#             raise HTTPException(
#                 status_code=500,
#                 detail="PDF was not created"
#             )

#         with open(pdf_path, "rb") as f:
#             pdf_bytes = f.read()

#         return Response(
#             content=pdf_bytes,
#             media_type="application/pdf",
#             headers={
#                 "Content-Disposition": f'attachment; filename="{pdf_filename}"'
#             }
#         )

#     except subprocess.TimeoutExpired:
#         raise HTTPException(status_code=408, detail="Conversion timed out")
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#     finally:
#         shutil.rmtree(tmp_dir, ignore_errors=True)


# @app.post("/convert/pdf-to-word")
# async def pdf_to_word(file: UploadFile = File(...)):
#     if not file.filename.endswith('.pdf'):
#         raise HTTPException(
#             status_code=400,
#             detail="Only .pdf files are supported"
#         )

#     tmp_dir = tempfile.mkdtemp()
#     try:
#         input_path = os.path.join(tmp_dir, file.filename)
#         with open(input_path, "wb") as f:
#             content = await file.read()
#             f.write(content)

#         result = subprocess.run([
#             "libreoffice",
#             "--headless",
#             "--convert-to", "docx:MS Word 2007 XML",
#             "--outdir", tmp_dir,
#             input_path
#         ], capture_output=True, text=True, timeout=60)

#         if result.returncode != 0:
#             raise HTTPException(
#                 status_code=500,
#                 detail=f"Conversion failed: {result.stderr}"
#             )

#         docx_filename = os.path.splitext(file.filename)[0] + ".docx"
#         docx_path = os.path.join(tmp_dir, docx_filename)

#         if not os.path.exists(docx_path):
#             raise HTTPException(
#                 status_code=500,
#                 detail="DOCX was not created"
#             )

#         with open(docx_path, "rb") as f:
#             docx_bytes = f.read()

#         return Response(
#             content=docx_bytes,
#             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             headers={
#                 "Content-Disposition": f'attachment; filename="{docx_filename}"'
#             }
#         )

#     except subprocess.TimeoutExpired:
#         raise HTTPException(status_code=408, detail="Conversion timed out")
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#     finally:
#         shutil.rmtree(tmp_dir, ignore_errors=True)







# from fastapi import FastAPI, UploadFile, File, HTTPException
# from fastapi.responses import Response
# import subprocess
# import tempfile
# import os
# import shutil

# app = FastAPI()

# @app.get("/health")
# def health():
#     return {"status": "ok"}

# # =====================================================================
# # ENDPOINT 1: WORD TO PDF (Kept exactly as you had it)
# # =====================================================================
# @app.post("/convert/word-to-pdf")
# async def word_to_pdf(file: UploadFile = File(...)):
#     if not file.filename.endswith(('.docx', '.doc')):
#         raise HTTPException(status_code=400, detail="Only .docx and .doc files are supported")
#     tmp_dir = tempfile.mkdtemp()
#     try:
#         input_path = os.path.join(tmp_dir, file.filename)
#         with open(input_path, "wb") as f:
#             content = await file.read()
#             f.write(content)
#         result = subprocess.run([
#             "libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmp_dir, input_path
#         ], capture_output=True, text=True, timeout=60)
#         if result.returncode != 0:
#             raise HTTPException(status_code=500, detail=f"Conversion failed: {result.stderr}")
#         pdf_filename = os.path.splitext(file.filename)[0] + ".pdf"
#         pdf_path = os.path.join(tmp_dir, pdf_filename)
#         if not os.path.exists(pdf_path):
#             raise HTTPException(status_code=500, detail="PDF was not created")
#         with open(pdf_path, "rb") as f:
#             pdf_bytes = f.read()
#         return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{pdf_filename}"'})
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#     finally:
#         shutil.rmtree(tmp_dir, ignore_errors=True)


# # =====================================================================
# # ENDPOINT 2: PDF TO WORD (DIAGNOSTIC SYSTEM)
# # =====================================================================
# @app.post("/convert/pdf-to-word")
# async def pdf_to_word(file: UploadFile = File(...)):
#     if not file.filename.endswith('.pdf'):
#         raise HTTPException(status_code=400, detail="Only .pdf files are supported")

#     # TEST FLAG: If you see this exact error in Postman, the deploy succeeded!
#     # If you still see "500: DOCX was not created", Railway is definitively serving old code.
#     try:
#         from pdf2docx import Converter
#     except ImportError as import_err:
#         raise HTTPException(
#             status_code=500, 
#             detail=f"[DIAGNOSTIC-ACTIVE] Code updated, but pdf2docx library is missing in your requirements file: {str(import_err)}"
#         )

#     tmp_dir = tempfile.mkdtemp()
#     try:
#         input_path = os.path.join(tmp_dir, file.filename)
#         with open(input_path, "wb") as f:
#             content = await file.read()
#             f.write(content)

#         docx_filename = os.path.splitext(file.filename)[0] + ".docx"
#         docx_path = os.path.join(tmp_dir, docx_filename)

#         # Run conversion
#         cv = Converter(input_path)
#         cv.convert(docx_path, start=0, end=None)
#         cv.close()

#         if not os.path.exists(docx_path):
#             raise HTTPException(
#                 status_code=500,
#                 detail="[DIAGNOSTIC-ACTIVE] Conversion completed but system could not locate output file."
#             )

#         with open(docx_path, "rb") as f:
#             docx_bytes = f.read()

#         return Response(
#             content=docx_bytes,
#             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             headers={"Content-Disposition": f'attachment; filename="{docx_filename}"'}
#         )

#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"[DIAGNOSTIC-ACTIVE] Engine crash: {str(e)}")
#     finally:
#         shutil.rmtree(tmp_dir, ignore_errors=True)





# from fastapi import FastAPI, UploadFile, File, HTTPException
# from fastapi.responses import Response
# import tempfile
# import os
# import shutil
# import subprocess

# app = FastAPI()

# @app.get("/health")
# def health():
#     return {"status": "ok"}

# # =====================================================================
# # ENDPOINT 1: WORD TO PDF
# # =====================================================================
# @app.post("/convert/word-to-pdf")
# async def word_to_pdf(file: UploadFile = File(...)):
#     if not file.filename.endswith(('.docx', '.doc')):
#         raise HTTPException(status_code=400, detail="Only .docx and .doc files are supported")
#     tmp_dir = tempfile.mkdtemp()
#     try:
#         input_path = os.path.join(tmp_dir, file.filename)
#         with open(input_path, "wb") as f:
#             f.write(await file.read())

#         result = subprocess.run([
#             "libreoffice", "--headless", "--convert-to", "pdf",
#             "--outdir", tmp_dir, input_path
#         ], capture_output=True, text=True, timeout=60)

#         if result.returncode != 0:
#             raise HTTPException(status_code=500, detail=f"Conversion failed: {result.stderr}")

#         pdf_filename = os.path.splitext(file.filename)[0] + ".pdf"
#         pdf_path = os.path.join(tmp_dir, pdf_filename)

#         if not os.path.exists(pdf_path):
#             raise HTTPException(status_code=500, detail="PDF was not created")

#         with open(pdf_path, "rb") as f:
#             pdf_bytes = f.read()

#         return Response(
#             content=pdf_bytes,
#             media_type="application/pdf",
#             headers={"Content-Disposition": f'attachment; filename="{pdf_filename}"'}
#         )
#     except subprocess.TimeoutExpired:
#         raise HTTPException(status_code=408, detail="Conversion timed out")
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#     finally:
#         shutil.rmtree(tmp_dir, ignore_errors=True)


# # =====================================================================
# # ENDPOINT 2: PDF TO WORD  (enhanced pdf2docx)
# # =====================================================================
# @app.post("/convert/pdf-to-word")
# async def pdf_to_word(file: UploadFile = File(...)):
#     if not file.filename.endswith('.pdf'):
#         raise HTTPException(status_code=400, detail="Only .pdf files are supported")

#     tmp_dir = tempfile.mkdtemp()
#     try:
#         input_path = os.path.join(tmp_dir, file.filename)
#         with open(input_path, "wb") as f:
#             f.write(await file.read())

#         docx_filename = os.path.splitext(file.filename)[0] + ".docx"
#         docx_path = os.path.join(tmp_dir, docx_filename)

#         from pdf2docx import Converter, parse

#         cv = Converter(input_path)

#         # --- Core layout settings ---
#         cv.convert(
#             docx_path,
#             start=0,
#             end=None,

#             # --- Text & spacing ---
#             # Merge nearby text fragments into one block
#             # Lower = more aggressive merging (good for scattered text)
#             connected_border_tolerance=3.0,

#             # How close lines must be to merge into one paragraph (pts)
#             line_overlap_threshold=0.9,

#             # Gap between lines to still count as same paragraph (pts)
#             line_break_free_space_ratio=0.3,

#             # Treat lines closer than this as same paragraph
#             line_separate_threshold=5.0,

#             # --- Tables ---
#             # Detect borderless/implicit tables
#             # True = try harder to find tables without visible borders
#             parse_lattice_table=True,
#             parse_stream_table=True,

#             # Minimum content needed to treat something as a table cell
#             # Lower = detect more tables
#             float_layout_tolerance=0.1,

#             # --- Images ---
#             # Extract and embed images in the docx
#             ignore_page_margin=False,

#             # --- Multi-column ---
#             # Detect multi-column layouts
#             # Higher = more tolerant of uneven columns
#             page_margin_factor=0.5,
#         )

#         cv.close()

#         if not os.path.exists(docx_path):
#             raise HTTPException(status_code=500, detail="DOCX was not created")

#         with open(docx_path, "rb") as f:
#             docx_bytes = f.read()

#         return Response(
#             content=docx_bytes,
#             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             headers={"Content-Disposition": f'attachment; filename="{docx_filename}"'}
#         )

#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")
#     finally:
#         shutil.rmtree(tmp_dir, ignore_errors=True)







from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import Response
import subprocess
import tempfile
import os
import shutil

app = FastAPI()

@app.get("/health")
def health():
    return {"status": "ok"}

# =====================================================================
# ENDPOINT 1: WORD TO PDF
# =====================================================================
@app.post("/convert/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    if not file.filename.endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Only .docx and .doc files are supported")
    tmp_dir = tempfile.mkdtemp()
    try:
        input_path = os.path.join(tmp_dir, file.filename)
        with open(input_path, "wb") as f:
            f.write(await file.read())
        result = subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", tmp_dir, input_path
        ], capture_output=True, text=True, timeout=60)
        if result.returncode != 0:
            raise HTTPException(status_code=500, detail=f"Conversion failed: {result.stderr}")
        pdf_filename = os.path.splitext(file.filename)[0] + ".pdf"
        pdf_path = os.path.join(tmp_dir, pdf_filename)
        if not os.path.exists(pdf_path):
            raise HTTPException(status_code=500, detail="PDF was not created")
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{pdf_filename}"'}
        )
    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=408, detail="Conversion timed out")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# =====================================================================
# ENDPOINT 2: PDF TO WORD  (pymupdf + python-docx)
# =====================================================================

import fitz  # pymupdf
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re
from dataclasses import dataclass, field
from typing import List, Optional, Tuple
from collections import defaultdict


@dataclass
class TextBlock:
    text: str
    x0: float
    y0: float
    x1: float
    y1: float
    font_size: float
    font_name: str
    is_bold: bool
    is_italic: bool
    color: Tuple[int, int, int]
    block_type: str = "text"   # text / image / table_cell
    spans: list = field(default_factory=list)


@dataclass
class TableData:
    rows: List[List[str]]
    x0: float
    y0: float
    x1: float
    y1: float


def _rgb_from_int(color_int: int) -> Tuple[int, int, int]:
    """Convert pymupdf integer color to (r, g, b) 0-255."""
    if color_int is None or color_int == 0:
        return (0, 0, 0)
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return (r, g, b)


def _is_bold(font_name: str) -> bool:
    name = font_name.lower()
    return any(k in name for k in ("bold", "-bd", ",bold", "black", "heavy", "demi"))


def _is_italic(font_name: str) -> bool:
    name = font_name.lower()
    return any(k in name for k in ("italic", "oblique", "-it", "-ob"))


def _classify_block(block: TextBlock, page_width: float, body_font_size: float) -> str:
    """
    Returns: heading1 / heading2 / heading3 / bullet / body
    """
    size = block.font_size
    text = block.text.strip()

    # Bullet detection
    if re.match(r'^[\u2022\u2023\u2043\u2219\-\*]\s+', text):
        return "bullet"
    if re.match(r'^\d+[\.\)]\s+', text):
        return "numbered"

    # Heading by font size relative to body
    if size >= body_font_size * 1.8 or (block.is_bold and size >= body_font_size * 1.5):
        return "heading1"
    if size >= body_font_size * 1.4 or (block.is_bold and size >= body_font_size * 1.25):
        return "heading2"
    if size >= body_font_size * 1.15 or (block.is_bold and size >= body_font_size * 1.1):
        return "heading3"

    return "body"


def _detect_body_font_size(blocks: List[TextBlock]) -> float:
    """Find the most common font size — that's the body size."""
    if not blocks:
        return 11.0
    size_counts = defaultdict(int)
    for b in blocks:
        size_counts[round(b.font_size, 1)] += len(b.text)
    return max(size_counts, key=size_counts.get)


def _blocks_overlap_vertically(b1: TextBlock, b2: TextBlock, threshold: float = 3.0) -> bool:
    return not (b1.y1 + threshold < b2.y0 or b2.y1 + threshold < b1.y0)


def _merge_line_blocks(blocks: List[TextBlock], page_width: float) -> List[TextBlock]:
    """
    Merge text blocks that are on the same line (same Y range) into one block,
    preserving left-to-right reading order. This fixes scattered text fragments.
    """
    if not blocks:
        return blocks

    # Sort by Y then X
    blocks = sorted(blocks, key=lambda b: (round(b.y0, 1), b.x0))
    merged = []
    used = [False] * len(blocks)

    for i, base in enumerate(blocks):
        if used[i]:
            continue
        group = [base]
        for j, other in enumerate(blocks):
            if i == j or used[j]:
                continue
            # Same line: Y overlap and similar font size
            y_overlap = not (base.y1 < other.y0 - 2 or other.y1 < base.y0 - 2)
            size_match = abs(base.font_size - other.font_size) < 2.0
            if y_overlap and size_match:
                group.append(other)
                used[j] = True
        used[i] = True

        # Sort group left to right
        group = sorted(group, key=lambda b: b.x0)

        # Merge spans and text
        merged_spans = []
        text_parts = []
        for g in group:
            merged_spans.extend(g.spans)
            text_parts.append(g.text.strip())

        # Join with space, avoid double spaces
        merged_text = " ".join(t for t in text_parts if t)

        merged_block = TextBlock(
            text=merged_text,
            x0=min(g.x0 for g in group),
            y0=min(g.y0 for g in group),
            x1=max(g.x1 for g in group),
            y1=max(g.y1 for g in group),
            font_size=base.font_size,
            font_name=base.font_name,
            is_bold=base.is_bold,
            is_italic=base.is_italic,
            color=base.color,
            spans=merged_spans,
        )
        merged.append(merged_block)

    return sorted(merged, key=lambda b: (round(b.y0, 1), b.x0))


def _group_into_paragraphs(blocks: List[TextBlock]) -> List[List[TextBlock]]:
    """
    Group line-merged blocks into paragraphs based on vertical gap.
    A new paragraph starts when the gap between lines exceeds 1.5x line height.
    """
    if not blocks:
        return []

    paragraphs = []
    current = [blocks[0]]

    for prev, curr in zip(blocks, blocks[1:]):
        line_height = prev.y1 - prev.y0
        gap = curr.y0 - prev.y1
        # New paragraph if gap > 1.5x line height or font size changes significantly
        if gap > line_height * 1.5 or abs(prev.font_size - curr.font_size) > 3:
            paragraphs.append(current)
            current = [curr]
        else:
            current.append(curr)

    paragraphs.append(current)
    return paragraphs


def _detect_columns(blocks: List[TextBlock], page_width: float) -> List[List[TextBlock]]:
    """
    Detect if page has multi-column layout.
    Returns blocks reordered for correct reading (top-to-bottom per column).
    """
    if not blocks:
        return [blocks]

    # Find a significant horizontal gap in the middle of the page
    mid = page_width / 2
    left_blocks = [b for b in blocks if b.x1 < mid + 20]
    right_blocks = [b for b in blocks if b.x0 > mid - 20]

    # If both columns have reasonable content, treat as two-column
    if (len(left_blocks) > 2 and len(right_blocks) > 2 and
            len(left_blocks) + len(right_blocks) >= len(blocks) * 0.85):
        left_sorted = sorted(left_blocks, key=lambda b: b.y0)
        right_sorted = sorted(right_blocks, key=lambda b: b.y0)
        return [left_sorted, right_sorted]

    return [sorted(blocks, key=lambda b: (b.y0, b.x0))]


def _detect_tables(page: fitz.Page) -> List[TableData]:
    """
    Use pymupdf's built-in table finder to detect tables on the page.
    Returns list of TableData with row/cell content.
    """
    tables = []
    try:
        tab_finder = page.find_tables()
        for tab in tab_finder.tables:
            rows = []
            for row in tab.extract():
                cells = [str(cell).strip() if cell else "" for cell in row]
                rows.append(cells)
            rect = tab.bbox
            tables.append(TableData(
                rows=rows,
                x0=rect[0], y0=rect[1],
                x1=rect[2], y1=rect[3],
            ))
    except Exception:
        pass
    return tables


def _rect_overlaps_table(block: TextBlock, tables: List[TableData]) -> bool:
    """Check if a text block falls inside a detected table area."""
    for t in tables:
        # Small margin to handle floating point
        if (block.x0 >= t.x0 - 5 and block.y0 >= t.y0 - 5 and
                block.x1 <= t.x1 + 5 and block.y1 <= t.y1 + 5):
            return True
    return False


def _add_table_to_doc(doc: Document, table_data: TableData):
    """Write a detected table into the Word document."""
    if not table_data.rows:
        return
    num_cols = max(len(row) for row in table_data.rows)
    if num_cols == 0:
        return

    tbl = doc.add_table(rows=len(table_data.rows), cols=num_cols)
    tbl.style = "Table Grid"

    for r_idx, row in enumerate(table_data.rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = tbl.cell(r_idx, c_idx)
                cell.text = cell_text
                # Bold first row as header
                if r_idx == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True


def _set_paragraph_format(para, block: TextBlock, block_type: str, doc_indent: float = 0):
    """Apply font, size, color, indent to a docx paragraph."""
    from docx.shared import Pt, RGBColor, Inches

    if block_type == "heading1":
        para.style = "Heading 1"
    elif block_type == "heading2":
        para.style = "Heading 2"
    elif block_type == "heading3":
        para.style = "Heading 3"
    else:
        para.style = "Normal"

    # Indent for bullets
    if block_type in ("bullet", "numbered"):
        para.paragraph_format.left_indent = Inches(0.25 * (1 + doc_indent))

    # Alignment
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.space_before = Pt(2)


def _add_runs_to_paragraph(para, block: TextBlock, block_type: str):
    """
    Add text runs to paragraph, preserving bold/italic/color per span.
    Falls back to block-level formatting if no spans.
    """
    from docx.shared import Pt, RGBColor

    spans = block.spans

    if not spans:
        run = para.add_run(block.text)
        run.bold = block.is_bold or block_type in ("heading1", "heading2", "heading3")
        run.italic = block.is_italic
        run.font.size = Pt(block.font_size)
        r, g, b = block.color
        if (r, g, b) != (0, 0, 0):
            run.font.color.rgb = RGBColor(r, g, b)
        return

    for span in spans:
        text = span.get("text", "").replace("\n", " ")
        if not text:
            continue
        run = para.add_run(text)
        fname = span.get("font", "")
        fsize = span.get("size", block.font_size)
        fcolor = span.get("color", 0)

        run.bold = _is_bold(fname) or block_type in ("heading1", "heading2", "heading3")
        run.italic = _is_italic(fname)
        run.font.size = Pt(fsize)

        r, g, b = _rgb_from_int(fcolor)
        if (r, g, b) != (0, 0, 0):
            run.font.color.rgb = RGBColor(r, g, b)


def _extract_page_blocks(page: fitz.Page) -> List[TextBlock]:
    """Extract all text blocks from a page with full span detail."""
    blocks = []
    raw = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

    for block in raw.get("blocks", []):
        if block.get("type") != 0:  # 0 = text
            continue

        all_spans = []
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                all_spans.append(span)

        if not all_spans:
            continue

        # Representative span = largest font size in block
        rep_span = max(all_spans, key=lambda s: s.get("size", 0))

        full_text = " ".join(
            span.get("text", "").strip()
            for line in block.get("lines", [])
            for span in line.get("spans", [])
            if span.get("text", "").strip()
        )

        if not full_text.strip():
            continue

        bbox = block.get("bbox", (0, 0, 0, 0))
        fname = rep_span.get("font", "")

        blocks.append(TextBlock(
            text=full_text,
            x0=bbox[0], y0=bbox[1],
            x1=bbox[2], y1=bbox[3],
            font_size=rep_span.get("size", 11),
            font_name=fname,
            is_bold=_is_bold(fname),
            is_italic=_is_italic(fname),
            color=_rgb_from_int(rep_span.get("color", 0)),
            spans=all_spans,
        ))

    return blocks


def convert_pdf_to_docx(input_path: str, output_path: str):
    pdf = fitz.open(input_path)
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        page_width = page.rect.width
        page_height = page.rect.height

        # Page break between pages
        if page_num > 0:
            doc.add_page_break()

        # ── 1. Detect tables first ────────────────────────────────────
        tables = _detect_tables(page)

        # ── 2. Extract images ─────────────────────────────────────────
        image_list = page.get_images(full=True)
        image_bboxes = []
        for img_info in image_list:
            xref = img_info[0]
            rects = page.get_image_rects(xref)
            for rect in rects:
                image_bboxes.append((xref, rect))

        # ── 3. Extract text blocks ────────────────────────────────────
        raw_blocks = _extract_page_blocks(page)

        # Filter out blocks inside table regions
        text_blocks = [
            b for b in raw_blocks
            if not _rect_overlaps_table(b, tables)
        ]

        # ── 4. Detect body font size ──────────────────────────────────
        body_size = _detect_body_font_size(text_blocks)

        # ── 5. Merge fragments on same line ──────────────────────────
        merged = _merge_line_blocks(text_blocks, page_width)

        # ── 6. Detect columns ────────────────────────────────────────
        columns = _detect_columns(merged, page_width)

        # ── 7. Group into paragraphs per column ──────────────────────
        all_paragraphs: List[List[TextBlock]] = []
        for col_blocks in columns:
            all_paragraphs.extend(_group_into_paragraphs(col_blocks))

        # Sort paragraph groups by Y position (top of first block)
        all_paragraphs.sort(key=lambda pg: pg[0].y0 if pg else 0)

        # Track which Y positions have tables/images already inserted
        inserted_tables = set()
        table_y_positions = [(t.y0, i) for i, t in enumerate(tables)]

        # ── 8. Write content in reading order ────────────────────────
        # Interleave tables and text by Y position
        content_items = []  # (y0, type, data)

        for i, t in enumerate(tables):
            content_items.append((t.y0, "table", t))

        for img_xref, img_rect in image_bboxes:
            content_items.append((img_rect.y0, "image", (img_xref, img_rect, page)))

        for para_group in all_paragraphs:
            if para_group:
                content_items.append((para_group[0].y0, "para", para_group))

        content_items.sort(key=lambda x: x[0])

        written_tables = set()
        written_images = set()

        for y0, ctype, data in content_items:

            if ctype == "table":
                t_id = id(data)
                if t_id not in written_tables:
                    written_tables.add(t_id)
                    _add_table_to_doc(doc, data)
                    doc.add_paragraph()  # spacing after table

            elif ctype == "image":
                img_xref, img_rect, pg = data
                img_id = img_xref
                if img_id not in written_images:
                    written_images.add(img_id)
                    try:
                        base_image = pdf.extract_image(img_xref)
                        img_bytes = base_image["image"]
                        img_ext = base_image["ext"]
                        img_stream = io.BytesIO(img_bytes)

                        # Scale image to fit page width
                        img_w = img_rect.width
                        img_h = img_rect.height
                        max_width = Inches(6.0)
                        aspect = img_h / img_w if img_w > 0 else 1
                        width = min(max_width, Inches(img_w / 72))
                        height = width * aspect

                        p = doc.add_paragraph()
                        run = p.add_run()
                        run.add_picture(img_stream, width=width)
                        p.paragraph_format.space_after = Pt(6)
                    except Exception:
                        pass

            elif ctype == "para":
                para_group = data
                if not para_group:
                    continue

                # Combine all blocks in paragraph group into one paragraph
                rep_block = max(para_group, key=lambda b: len(b.text))
                block_type = _classify_block(rep_block, page_width, body_size)

                # Combine text from all blocks in group
                combined_spans = []
                for b in para_group:
                    combined_spans.extend(b.spans)

                combined_block = TextBlock(
                    text=" ".join(b.text.strip() for b in para_group),
                    x0=min(b.x0 for b in para_group),
                    y0=para_group[0].y0,
                    x1=max(b.x1 for b in para_group),
                    y1=para_group[-1].y1,
                    font_size=rep_block.font_size,
                    font_name=rep_block.font_name,
                    is_bold=rep_block.is_bold,
                    is_italic=rep_block.is_italic,
                    color=rep_block.color,
                    spans=combined_spans,
                )

                para = doc.add_paragraph()
                _set_paragraph_format(para, combined_block, block_type)
                _add_runs_to_paragraph(para, combined_block, block_type)

    pdf.close()
    doc.save(output_path)


@app.post("/convert/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only .pdf files are supported")

    tmp_dir = tempfile.mkdtemp()
    try:
        input_path = os.path.join(tmp_dir, file.filename)
        with open(input_path, "wb") as f:
            f.write(await file.read())

        docx_filename = os.path.splitext(file.filename)[0] + ".docx"
        docx_path = os.path.join(tmp_dir, docx_filename)

        convert_pdf_to_docx(input_path, docx_path)

        if not os.path.exists(docx_path):
            raise HTTPException(status_code=500, detail="DOCX was not created")

        with open(docx_path, "rb") as f:
            docx_bytes = f.read()

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{docx_filename}"'}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)